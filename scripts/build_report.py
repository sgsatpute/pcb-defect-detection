from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
RESULTS_DIR = ROOT / "results" / "pcb_results"
ASSETS_DIR = ROOT / "assets"
REPORTS_DIR = ROOT / "reports"
OUTPUT = REPORTS_DIR / "PCB_Defect_Detection_Research_Paper_Improved.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(0, 0, 0)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"
BORDER = "D9E2F3"


def set_run_font(run, *, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border_bottom(paragraph, color="2E74B5", size="12", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def style_table(table, widths_dxa, header=True):
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    if header:
        mark_header_row(table.rows[0])
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if header and row_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=10,
                        color=INK,
                        bold=(header and row_idx == 0),
                    )


def add_caption(doc, text):
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(text)
    set_run_font(run, size=9, italic=True, color=MUTED)


def add_body_paragraph(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.333
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=INK)
    return paragraph


def add_heading(doc, text, level):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    if level == 1:
        paragraph.paragraph_format.space_before = Pt(18)
        paragraph.paragraph_format.space_after = Pt(10)
        size, color = 16, BLUE
    elif level == 2:
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        size, color = 13, BLUE
    else:
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
        size, color = 12, DARK_BLUE
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=True)
    return paragraph


def add_picture(doc, image_path, caption, max_width=6.15):
    with Image.open(image_path) as image:
        width_px, height_px = image.size
    width = min(max_width, width_px / 300)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(image_path), width=Inches(width))
    inline_shape._inline.docPr.set("descr", caption)
    inline_shape._inline.docPr.set("title", image_path.name)
    add_caption(doc, caption)


def load_metrics():
    rows = []
    with (RESULTS_DIR / "results.csv").open(newline="") as file:
        for row in csv.DictReader(file):
            rows.append({key.strip(): float(value) for key, value in row.items() if value != ""})
    best = {
        "Precision": max(rows, key=lambda row: row["metrics/precision(B)"]),
        "Recall": max(rows, key=lambda row: row["metrics/recall(B)"]),
        "mAP@0.5": max(rows, key=lambda row: row["metrics/mAP50(B)"]),
        "mAP@0.5:0.95": max(rows, key=lambda row: row["metrics/mAP50-95(B)"]),
    }
    return rows, best, rows[-1]


def build_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            row.cells[idx].text = str(value)
    style_table(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.bold = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("PCB Defect Detection - YOLOv8s")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("COEP Technological University | DeepPCB automated visual inspection")
    set_run_font(run, size=9, color=MUTED)


def build_report():
    rows, best, last = load_metrics()
    doc = Document()
    configure_document(doc)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("Research Report")
    set_run_font(run, size=11, color=DARK_BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("PCB Defect Detection Using YOLOv8s")
    set_run_font(run, size=24, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("A Deep Learning Pipeline for Automated Visual Inspection")
    set_run_font(run, size=13, color=MUTED)

    metadata = [
        ("Author", "Saurav Satpute"),
        ("Institution", "Department of Electronics and Telecommunication Engineering, COEP Technological University"),
        ("Artifacts", "YOLOv8s checkpoint, Streamlit app, notebook, training outputs, and qualitative result image"),
    ]
    for label_text, value_text in metadata:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(2)
        label = paragraph.add_run(f"{label_text}: ")
        set_run_font(label, size=10.5, color=INK, bold=True)
        value = paragraph.add_run(value_text)
        set_run_font(value, size=10.5, color=MUTED)

    rule = doc.add_paragraph()
    set_paragraph_border_bottom(rule)
    rule.paragraph_format.space_after = Pt(14)

    add_heading(doc, "Abstract", 1)
    add_body_paragraph(
        doc,
        "Printed Circuit Board manufacturing is vulnerable to small defects such as open traces, shorts, "
        "pin-holes, spurs, mousebites, and residual copper. These defects can reduce electrical reliability "
        "and are difficult to inspect manually at production scale. This report presents a YOLOv8s-based "
        "detector trained on the DeepPCB dataset and packaged with a Streamlit demonstration application. "
        "The logged training run achieved a best mAP@0.5 of 0.99376 and a best mAP@0.5:0.95 of 0.78211, "
        "showing strong class-level localization at standard IoU while leaving room for tighter-box "
        "improvement under stricter IoU thresholds.",
    )

    keywords = doc.add_paragraph()
    keywords.paragraph_format.space_after = Pt(10)
    label = keywords.add_run("Keywords: ")
    set_run_font(label, size=10.5, color=INK, bold=True)
    value = keywords.add_run("PCB defect detection, YOLOv8, DeepPCB, object detection, automated visual inspection")
    set_run_font(value, size=10.5, color=MUTED)

    add_heading(doc, "1. Introduction", 1)
    add_body_paragraph(
        doc,
        "PCBs are foundational to modern electronic assemblies, and even sub-millimetre manufacturing "
        "defects can interrupt signal paths, create unintended bridges, or degrade long-term reliability. "
        "Traditional template matching and morphology-based inspection pipelines work well for stable "
        "production lines but require frequent retuning when board layouts, lighting, or camera setups change.",
    )
    add_body_paragraph(
        doc,
        "Single-stage object detectors such as YOLO are well suited to this problem because they combine "
        "localization and classification in one pass. The objective of this work is to train a compact "
        "YOLOv8s model for six DeepPCB defect classes and package the trained model in a simple interactive "
        "application that can demonstrate inference on uploaded PCB images.",
    )

    add_heading(doc, "2. Dataset and Defect Taxonomy", 1)
    add_body_paragraph(
        doc,
        "The DeepPCB dataset contains paired PCB images with annotated defect regions. The project converts "
        "the original bounding boxes into YOLO format and trains the detector at 640 x 640 input resolution. "
        "The six-class taxonomy used by the model is summarized below.",
    )
    build_table(
        doc,
        ["Class", "Manufacturing meaning"],
        [
            ["open", "Missing copper trace or unintended circuit break."],
            ["short", "Unintended copper bridge between traces that should remain isolated."],
            ["mousebite", "Small notch removed from a trace edge, reducing copper cross-section."],
            ["spur", "Excess copper protrusion extending from a trace edge."],
            ["copper", "Residual copper spot that should have been etched away."],
            ["pin-hole", "Small hole in a copper region caused by incomplete plating or etching."],
        ],
        [1900, 7460],
    )

    add_heading(doc, "3. Methodology", 1)
    add_body_paragraph(
        doc,
        "YOLOv8s was selected because it offers a practical balance between inference speed and detection "
        "accuracy. The model uses a CSP-style feature extraction backbone, multi-scale feature aggregation, "
        "and an anchor-free detection head. Transfer learning from the pretrained YOLOv8s checkpoint reduces "
        "training time and improves convergence on the industrial inspection task.",
    )
    build_table(
        doc,
        ["Training parameter", "Value from run artifact"],
        [
            ["Base model", "yolov8s.pt"],
            ["Task", "Object detection"],
            ["Image size", "640 x 640 px"],
            ["Batch size", "16"],
            ["Configured epochs", "50"],
            ["Early stopping patience", "10 epochs"],
            ["Logged epochs", f"{len(rows)}"],
            ["Last logged training time", f"{last['time'] / 60:.1f} minutes"],
            ["Augmentation", "HSV jitter, translation, scale, horizontal flip, mosaic, erasing"],
            ["Validation split", "DeepPCB validation subset"],
        ],
        [2600, 6760],
    )

    add_heading(doc, "4. Quantitative Results", 1)
    add_body_paragraph(
        doc,
        "The training log shows fast convergence. The model exceeded mAP@0.5 of 0.95 by epoch 4, passed "
        "0.98 by epoch 13, and reached its best logged mAP@0.5 at epoch 32. The stricter mAP@0.5:0.95 "
        "metric peaked earlier, which indicates that standard defect classification is strong while "
        "fine-grained box tightness remains the main area for further improvement.",
    )
    build_table(
        doc,
        ["Metric", "Best epoch", "Best value", "Final logged value"],
        [
            ["Precision", int(best["Precision"]["epoch"]), f"{best['Precision']['metrics/precision(B)']:.5f}", f"{last['metrics/precision(B)']:.5f}"],
            ["Recall", int(best["Recall"]["epoch"]), f"{best['Recall']['metrics/recall(B)']:.5f}", f"{last['metrics/recall(B)']:.5f}"],
            ["mAP@0.5", int(best["mAP@0.5"]["epoch"]), f"{best['mAP@0.5']['metrics/mAP50(B)']:.5f}", f"{last['metrics/mAP50(B)']:.5f}"],
            ["mAP@0.5:0.95", int(best["mAP@0.5:0.95"]["epoch"]), f"{best['mAP@0.5:0.95']['metrics/mAP50-95(B)']:.5f}", f"{last['metrics/mAP50-95(B)']:.5f}"],
        ],
        [2500, 1900, 2480, 2480],
    )

    add_picture(
        doc,
        RESULTS_DIR / "results.png",
        "Figure 1. Training losses and validation metrics from the YOLOv8s run.",
        max_width=6.15,
    )
    add_picture(
        doc,
        RESULTS_DIR / "BoxPR_curve.png",
        "Figure 2. Precision-recall curve for bounding-box detection.",
        max_width=5.9,
    )
    add_picture(
        doc,
        RESULTS_DIR / "BoxF1_curve.png",
        "Figure 3. F1 score as a function of confidence threshold.",
        max_width=5.9,
    )

    add_heading(doc, "5. Confusion Matrix and Error Profile", 1)
    add_body_paragraph(
        doc,
        "The normalized confusion matrix is dominated by diagonal cells, indicating that most errors are "
        "localization or confidence-threshold issues rather than systematic confusion between defect classes. "
        "For deployment, low-confidence detections should be reviewed carefully because tiny pin-hole and "
        "mousebite defects are visually similar to noise or imperfect segmentation boundaries.",
    )
    add_picture(
        doc,
        RESULTS_DIR / "confusion_matrix_normalized.png",
        "Figure 4. Normalized validation confusion matrix.",
        max_width=5.6,
    )

    add_heading(doc, "6. Qualitative Inference Results", 1)
    add_body_paragraph(
        doc,
        "The included qualitative result image demonstrates multi-defect detection across four PCB crops. "
        "The model identifies all six supported classes, including small pin-hole defects, edge notches, "
        "spurs, residual copper, open traces, and shorts. The sample contains visible counts of 6, 5, 8, "
        "and 6 defects across the four panels, for 25 displayed detections in total.",
    )
    add_picture(
        doc,
        ASSETS_DIR / "detection_results.png",
        "Figure 5. Example YOLOv8s detections on PCB samples.",
        max_width=5.55,
    )

    add_heading(doc, "7. Streamlit Demonstration Application", 1)
    add_body_paragraph(
        doc,
        "The project includes a Streamlit application that loads the packaged best.pt checkpoint, accepts "
        "uploaded PCB images, runs YOLO inference, displays the annotated result, and summarizes detections "
        "by class with maximum and average confidence values. The application also exposes confidence and "
        "NMS IOU controls, allowing users to trade off sensitivity and false-positive filtering without "
        "editing code.",
    )

    build_table(
        doc,
        ["Application component", "Purpose"],
        [
            ["Model selector", "Loads the local best.pt checkpoint or an uploaded YOLO .pt model."],
            ["Image uploader", "Accepts JPG, JPEG, PNG, and BMP PCB images."],
            ["Threshold controls", "Adjusts confidence and IOU values for inference-time filtering."],
            ["Detection breakdown", "Counts defect types and reports confidence summaries."],
            ["Download button", "Exports the annotated image for documentation or review."],
        ],
        [2700, 6660],
    )

    add_heading(doc, "8. Limitations and Future Work", 1)
    add_body_paragraph(
        doc,
        "The strongest current result is high mAP@0.5, but the lower and more variable mAP@0.5:0.95 shows "
        "that tighter localization remains the most important technical improvement. Practical deployment "
        "would also require validation under camera noise, production lighting variation, unseen PCB layouts, "
        "and real-time throughput constraints.",
    )
    build_table(
        doc,
        ["Future direction", "Expected benefit"],
        [
            ["TensorRT or OpenVINO export", "Lower inference latency for edge deployment."],
            ["YOLOv8n distillation", "Smaller model for low-power inspection stations."],
            ["Synthetic augmentation", "Better generalization to unseen PCB geometries and lighting."],
            ["Active learning loop", "Prioritize human review of low-confidence or novel defects."],
            ["Anomaly detection layer", "Flag defect types not present in the six-class training taxonomy."],
        ],
        [2800, 6560],
    )

    add_heading(doc, "9. Conclusion", 1)
    add_body_paragraph(
        doc,
        "The YOLOv8s PCB defect detector provides a strong baseline for automated visual inspection on "
        "the DeepPCB task. The packaged repository contains the trained checkpoint, training outputs, a "
        "clean training notebook, a Streamlit inference application, and reproducible documentation. The "
        "best logged mAP@0.5 of 0.99376 confirms that the model can reliably recognize the six supported "
        "defect classes under the benchmark conditions, while stricter IoU metrics identify localization "
        "precision as the key next step.",
    )

    add_heading(doc, "References", 1)
    references = [
        '[1] S. Tang, F. He, X. Huang, and J. Yang, "Online PCB defect detector on a new PCB defect dataset," arXiv:1902.06197, 2019.',
        '[2] W. Huang and P. Wei, "A PCB dataset for defects detection and classification," arXiv:1901.08204, 2019.',
        '[3] G. Jocher, A. Chaurasia, and J. Qiu, "Ultralytics YOLO," 2023. https://github.com/ultralytics/ultralytics',
        '[4] J. Redmon, S. Divvala, R. Girshick, and A. Farhadi, "You Only Look Once: Unified, real-time object detection," CVPR, 2016.',
        '[5] S. Ren, K. He, R. Girshick, and J. Sun, "Faster R-CNN: Towards real-time object detection with region proposal networks," IEEE TPAMI, 2017.',
    ]
    for reference in references:
        add_body_paragraph(doc, reference)

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
